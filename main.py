from pytesseract import pytesseract
from os.path import isfile, join
from docx import Document
from os import listdir
import sys
import cv2
import os

class ImageConversor:
    def __init__(self):

        # Argument counter
        self.argc = len(sys.argv)

        # Files or arguments
        self.mode = None

        # Images
        self.images = []

        # String to convert
        self.paragraphs = []

        # https://github.com/UB-Mannheim/tesseract/wiki
        
        # tesseract path
        pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


    def run (self):
        '''Runs the main functionality'''

        # Checks the arguments provided
        self.check_args()

        # Stores images in self.images
        self.get_images()

        print(f"Images: {self.images}")

        if len(self.images) < 1:
            print("No images were found")
            exit()

        for img in self.images:
            self.image_processor(img)

        self.text_converter()

        
    def text_converter(self):
        '''Converts the text to a Word Document'''

        # Initializes the document 
        document = Document()

        # Adds a heading
        document.add_heading('Document Auto Generated with Python', 0)

        # Adds all the paragraphs
        for p in self.paragraphs:
            document.add_paragraph(p)
        
        # Saves the document
        document.save('document.docx')
        

    def image_processor(self, filename):
        '''Gets the text of the image '''
        # Reads image with opencv
        image = cv2.imread(filename)

        # Converts image to string
        data = pytesseract.image_to_string(image)

        # 
        self.paragraphs.append(data)

    def get_images(self):
        '''Stores in self.images all the images provided'''

        # If user is using a directory
        if self.mode == "d":
            # Checks if the directory exists
            if os.path.exists(sys.argv[2]):
                # Saves files with the extension provided
                self.images = [join(sys.argv[2], f) for f in listdir(sys.argv[2]) 
                               if isfile(join(sys.argv[2], f)) and
                               f.endswith(sys.argv[3])]
            else:
                print("Directory does not exists")
                exit()
        else:
            for f in range(2, len(sys.argv)):
                # Adds every image that is a file
                if isfile(sys.argv[f]):
                    self.images.append(sys.argv[f])

    def check_args(self):
        '''Checks users parameters'''
        if self.argc < 3:
            # Not enough arguments were provided
            print("Error. Expecting at least 2 argument.")
            exit()
        elif sys.argv[1] == "-f":
            # User is providing files
            self.mode = "f"
        elif sys.argv[1] == "-d":
            if self.argc != 4:
                print("You must include directory and extension.")
                exit()
            # User is providing directories
            self.mode = "d"
        elif sys.argv[1] == "-h":
            # Show help
            pass
        else:
            print("Please, select a mode")


if __name__ == '__main__':
    ImageConversor().run()
